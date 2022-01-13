import pandas as pd
import pandas.io.formats.excel
from django.apps import apps
from django.db.backends.mysql.base import DatabaseWrapper
from django.db.models.fields import NOT_PROVIDED
from tqdm import tqdm
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor
from docx.text.parfmt import ParagraphFormat
from docx.oxml import parse_xml

pandas.io.formats.excel.header_style = None
data_types = DatabaseWrapper.data_types


def get_models(model_filter=lambda m: True):
    models = apps.get_models()
    if model_filter and callable(model_filter):
        models = [model for model in models if model_filter(model)]
    return models


def collect_table_info(models):
    df = pd.DataFrame()
    print('collect infos....')
    for model in tqdm(models):
        _meta = model._meta
        table_name = _meta.original_attrs.get('db_table', _meta.label_lower.replace('.', '_'))
        for field in _meta.fields:
            default = field.default
            if default == NOT_PROVIDED:
                default = None
            elif callable(default):
                try:
                    default = default.__name__
                except AttributeError:
                    pass

            df = df.append({
                '表名': table_name,
                '名称': field.name,
                '主键': field.primary_key,
                '类型': data_types.get(field.get_internal_type(), '') % field.__dict__,
                '非空约束': field.null,
                '默认值': default,
                '唯一约束': field.unique,
                '说明': str(field.verbose_name)
            }, ignore_index=True)
    return df


def write_to_excel(df=pd.DataFrame()):
    print('write to excel....')
    df.set_index(['表名', '名称'], inplace=True)
    df.to_excel("数据库说明文档.xlsx")


def write_to_doc(df=pd.DataFrame()):
    print('write doc....')
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = u'宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    style._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    document.add_heading('数据库设计说明书', 0)

    # table list
    document.add_heading(u'1.数据库表名列表', 1)
    table = document.add_table(rows=1, cols=3, style='TableGrid')
    table.style.paragraph_format.space_after = Inches(0)
    table.style.font.size = Pt(10)

    for i, (col_name, col_width) in enumerate(zip(['序号', '表名', '说明'], [1, 20, 20])):
        cell = table.rows[0].cells[i]
        cell.width = Inches(col_width)
        cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="9CC2E5"/>'.format(nsdecls('w'))))
        cell.paragraphs[0].add_run(col_name)

    table_list = sorted(set(df['表名'].values.tolist()))
    for i, t in enumerate(table_list):
        cells = table.add_row().cells
        cells[0].paragraphs[0].add_run(str(i + 1))
        cells[1].paragraphs[0].add_run(t)

    # table detail
    document.add_page_break()
    document.add_heading(u'2.数据库详细设计', 1)
    columns = ["名称", "类型", "主键", "默认值", "非空约束", "唯一约束", "说明"]
    df.set_index('表名', inplace=True)

    for t in tqdm(table_list):
        document.add_paragraph(t, style='List Bullet').paragraph_format.space_after = Inches(0)
        table = document.add_table(rows=1, cols=7, style='TableGrid')
        table.style.paragraph_format.space_after = Inches(0)
        table.style.font.size = Pt(10)
        cells = table.rows[0].cells

        for i, col in enumerate(columns):
            cells[i]._tc.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="9CC2E5"/>'.format(nsdecls('w'))))
            cells[i].paragraphs[0].add_run(col)

        for _, df_row in df.loc[t, :].iterrows():
            row = table.add_row()
            for i, cell in enumerate(row.cells):
                text = df_row[columns[i]]
                if isinstance(text, bool):
                    text = '✓' if text else ''
                elif not text:
                    text = ''
                else:
                    text = str(text)
                cell.paragraphs[0].add_run(text)
        document.add_paragraph()

    document.save('数据库说明文档.docx')


def run(*args, **kwargs):
    target = args[0] if args else 'doc'
    if target == 'doc':
        write_to_doc(collect_table_info(get_models()))
    else:
        write_to_excel(collect_table_info(get_models()))




